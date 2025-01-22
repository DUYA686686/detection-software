import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_LINE_SPACING

fileList = []  # 使用全局列表保存文件路径


def getAllFile(path, fileList):  # 使用递归方法
    dirList = []  # 保存文件夹
    files = os.listdir(path)  # 返回一个列表，其中包含文件和文件夹
    if '.DS_Store' in files:
        files.remove('.DS_Store')
    for f in files:
        print(f)
        if (os.path.isdir(path + '/' + f)):
            dirList.append(path + '/' + f)  # 将文件夹名字进行保存
        if (os.path.isfile(path + '/' + f)):
            fileList.append(path + '/' + f)  # 将文件名保存
    for dir in dirList:  # 如果文件夹为空时，递归自动退出
        getAllFile(dir, fileList)  # 递归保存到将所有文件保存到fileList中


def getExpectedFile(fileList):
    for file in fileList:
        if not file.endswith('.py'):  # 删除不是所期望后缀文件的格式
            fileList.remove(file)
    print('python文件数量为： ', len(fileList))


def saveDocFile():
    doc = Document()
    p = doc.add_paragraph('')  # 增加一页
    doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
    doc.styles['header'].font.name = u'黑体'
    doc.styles['Normal'].font.size = Pt(9)  # 设置小五字体
    p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = Pt(12.9)  # 固定值12.9磅, 保证每页有50行代码
    save_file = r'基于深度学习Faster R-IR7-EC的桥梁表面缺陷识别.doc'
    codeNum = 0

    for i, f in enumerate(fileList):
        print('starting deal %d' % i)

        with open(f, encoding='UTF-8') as file:  # 转换编码以实现正确输出中文格式
            for line in file.readlines():
                if line == '\n':  # 删除空行
                    continue
                if re.match(r'^\s+$', line):  # 使用正则表达式删除全是空格的空行
                    continue

                p.add_run(line)
                codeNum += 1  # 记录是已经写入的数据
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                paragraph.text = '基于深度学习Faster R-IR7-EC的桥梁表面缺陷识别软件V1.0\t\t'  # 在页眉添加内容，制表符分隔左、中、右
                paragraph.style = doc.styles['header']

                if codeNum == 3050:  # 保证小于等于60页
                    doc.save(save_file)
                    return
    doc.save(save_file)  # 不足60页进行保存


getAllFile(r'./软著封装', fileList)
print(fileList)
print('文件数量为： ', len(fileList))

##如果要合并所有后缀文件，则注释以下两行代码
# getExpectedFile(fileList)
# print(os.path.isfile(fileList[0])) # 判断第一个值是否是文件

saveDocFile()
print('完成源代码合并！')
